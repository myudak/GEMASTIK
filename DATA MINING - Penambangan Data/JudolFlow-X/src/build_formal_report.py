from __future__ import annotations

import json
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Iterable, List

import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
TABLE_DIR = REPORT_DIR / "tables"
FIGURE_DIR = REPORT_DIR / "figures"

DOCX_PATH = REPORT_DIR / "GEMASTIK XVIII Penambangan Data - ID_TIM - NAMA_TIM - JudolFlow-X.docx"
PDF_PATH = REPORT_DIR / "GEMASTIK XVIII Penambangan Data - ID_TIM - NAMA_TIM - JudolFlow-X.pdf"


TITLE = "JudolFlow-X: Hybrid NLP, Entity Mining, and Graph Ranking untuk Prioritas Tinjauan Promosi Judi Online"
SUBTITLE = "Technical Report GEMASTIK XVIII - Penambangan Data"


def load_table(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLE_DIR / name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_paragraph(doc: Document, text: str, style: str = "Body Text", bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    set_keep_with_next(p)
    return p


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.15)
    section.right_margin = Cm(2.15)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    body = styles["Body Text"]
    body.font.name = "Times New Roman"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    body.font.size = Pt(11)
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.08

    for style_name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(16, 24, 32)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(SUBTITLE)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(71, 85, 105)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(127, 29, 29)

    meta = [
        "ID Tim: ID_TIM",
        "Nama Tim: NAMA_TIM",
        "Perguruan Tinggi: NAMA_UNIVERSITAS",
        "Anggota: NAMA_ANGGOTA_1, NAMA_ANGGOTA_2, NAMA_ANGGOTA_3",
        f"Tanggal penyusunan artefak: {datetime.now().strftime('%d %B %Y')}",
    ]
    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(11)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "FFF7ED")
    set_cell_margins(cell, 170, 220, 170, 220)
    text = (
        "Batas etis: JudolFlow-X menghasilkan antrian tinjauan manusia terhadap cluster publik yang telah dimasking. "
        "Sistem tidak menuduh individu, tidak mempublikasikan domain/handle live, dan tidak menggantikan proses verifikasi lembaga berwenang."
    )
    cell.paragraphs[0].add_run(text).font.size = Pt(10)
    doc.add_page_break()


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    return p


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.9) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc: Document, df: pd.DataFrame, caption: str, widths: List[float] | None = None, max_rows: int | None = None) -> None:
    if max_rows is not None:
        df = df.head(max_rows)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(9)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        set_cell_shading(hdr[i], "E7E5DF")
        set_cell_margins(hdr[i])
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8.5)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            text = str(value)
            if len(text) > 160:
                text = text[:157] + "..."
            cells[i].text = text
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths[: len(row.cells)]):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def write_report() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_cover(doc)

    metrics = load_table("metrics_summary.csv")
    ranking = load_table("ranking_snapshot.csv")
    dataset = load_table("dataset_summary.csv")
    sources = load_table("source_status.csv")
    ablation = load_table("feature_ablation.csv")
    importance = load_table("feature_importance.csv")
    stability = load_table("ranking_stability.csv")
    cases = load_table("case_studies.csv")
    validation = load_table("validation_checks.csv")

    add_heading(doc, "Abstrak", 1)
    add_paragraph(
        doc,
        "Judi online menjadi masalah ruang digital karena promosi, tautan mirror, akun perantara, dan petunjuk pembayaran dapat muncul secara cepat serta lintas platform. "
        "JudolFlow-X mengusulkan pipeline penambangan data untuk memprioritaskan cluster promosi judi online yang perlu ditinjau oleh manusia. "
        "Unit analisis bukan individu, melainkan cluster konten publik dan entitas yang telah dimasking. Metode yang digunakan mencakup normalisasi teks Indonesia, ekstraksi entitas domain/handle/kontak/pembayaran, pembentukan graf ko-kemunculan, klasifikasi berbasis weak supervision, dan ranking multi-sinyal. "
        "Pada corpus prototipe anonim sebanyak 620 sampel, model StumpBoost JudolFlow memperoleh PR-AUC tinggi pada split waktu dan menghasilkan antrian Top-K yang stabil. "
        "Kontribusi utama karya ini adalah desain sistem yang dapat diuji ulang, aman secara etis, dan berorientasi pada prioritas tinjauan, bukan tuduhan otomatis.",
    )
    add_paragraph(doc, "Kata kunci: judi online, graph mining, NLP, weak supervision, entity extraction, ranking.")

    add_heading(doc, "1. Pendahuluan", 1)
    add_paragraph(
        doc,
        "Promosi judi online sering memanfaatkan bahasa tersandi, komentar berulang, tautan berganti, akun kontak, serta petunjuk pembayaran. "
        "Dalam skala besar, moderator atau lembaga pengawas menghadapi keterbatasan kapasitas: semua konten tidak mungkin diperiksa secara manual dengan kedalaman yang sama. "
        "Oleh karena itu, masalah data mining yang relevan bukan hanya klasifikasi satu konten, melainkan prioritisasi cluster yang paling layak masuk antrian tinjauan."
    )
    add_paragraph(
        doc,
        "Konteks nasional menunjukkan isu ini penting. Komdigi melaporkan pemutusan akses konten terkait judi online dalam skala jutaan, pengembangan SAMAN untuk pemantauan ruang digital, serta penyediaan kanal pelaporan publik seperti AduanKonten dan CekRekening. "
        "PPATK dan OJK juga menempatkan judi online sebagai risiko sosial-ekonomi dan finansial yang perlu ditangani lintas lembaga. "
        "JudolFlow-X berada pada lapisan pendukung keputusan: membantu menentukan cluster mana yang perlu dilihat lebih dahulu."
    )
    add_paragraph(
        doc,
        "Tujuan penelitian ini adalah membangun prototipe yang dapat: (1) mengekstraksi sinyal teks dan entitas dari konten publik, (2) membentuk graf ko-kemunculan entitas, (3) menghitung skor prioritas tinjauan cluster, dan (4) menjelaskan alasan ranking secara ringkas. "
        "Manfaat praktisnya adalah efisiensi triase, konsistensi pelaporan, dan transparansi fitur yang mendukung pemeriksaan manusia."
    )
    add_paragraph(
        doc,
        "Batasan penelitian: dataset bawaan adalah corpus anonim sintetik bergaya sinyal publik agar tidak menyebarkan domain, handle, nomor telepon, atau petunjuk pembayaran yang nyata. "
        "Untuk pengajuan final, tim perlu mengganti atau melengkapi corpus ini dengan snapshot publik yang diperoleh secara legal, dianotasi, dan disamarkan. "
        "Hasil model tidak boleh dipakai sebagai dasar tuduhan otomatis terhadap orang, akun, atau institusi."
    )

    add_heading(doc, "2. Kajian Terkait", 1)
    add_paragraph(
        doc,
        "Penelitian deteksi konten ilegal dan berbahaya umumnya memadukan NLP, ekstraksi entitas, dan analisis graf. "
        "Model teks mampu menangkap pola semantik seperti ajakan, promosi, dan instruksi transaksi, tetapi mudah gagal jika aktor menggunakan slang atau obfuscation. "
        "Sebaliknya, graf ko-kemunculan entitas dapat memperlihatkan pola koordinasi melalui domain, handle, kontak, dan petunjuk pembayaran yang dipakai berulang."
    )
    add_paragraph(
        doc,
        "Pendekatan retrieval dan reranking modern dapat digunakan sebagai tahap lanjutan. Embedding multilingual, dense retrieval, dan cross-encoder reranker berguna untuk menemukan variasi konten yang secara leksikal berbeda tetapi bermakna serupa. "
        "Namun, agar prototipe GEMASTIK dapat direplikasi tanpa GPU, implementasi awal menggunakan model tabular ringan dan graf connected components. Desain ini tetap membuka jalur pengembangan menuju BGE-M3, Qwen embedding, GraphSAGE, atau heterogeneous GNN ketika data berlabel lebih besar tersedia."
    )
    add_paragraph(
        doc,
        "Perbedaan JudolFlow-X dibanding klasifikasi konten sederhana adalah keluaran utamanya berupa ranking cluster. "
        "Ranking cluster lebih sesuai untuk kebutuhan operasional karena satu domain, handle, atau pola pembayaran dapat muncul dalam banyak sampel. "
        "Dengan demikian, evaluator dapat melihat konsentrasi bukti, bukan hanya satu teks terpisah."
    )

    add_heading(doc, "3. Dataset dan Formulasi Masalah", 1)
    add_paragraph(
        doc,
        "Unit data adalah sampel konten publik yang memiliki metadata platform, waktu, sumber, teks, dan label tinjauan. "
        "Pipeline mengekstraksi empat keluarga entitas: domain, handle, kontak telepon, dan petunjuk pembayaran. Seluruh entitas dimasking menjadi hash bertipe, misalnya domain_xxx atau handle_xxx, sebelum masuk tabel proses."
    )
    add_table(doc, dataset, "Tabel 1. Ringkasan dataset dan status masking.", widths=[2.4, 3.8])
    add_table(
        doc,
        validation,
        "Tabel 2. Validasi dataset sebelum eksperimen.",
        widths=[3.3, 1.6],
    )
    add_paragraph(
        doc,
        "Target utama adalah skor prioritas tinjauan cluster. Label konten menggunakan empat kelas operasional: benign, suspicious, promotional, dan unknown. "
        "Untuk eksperimen awal, suspicious dan promotional dianggap positif pada tugas klasifikasi konten. Untuk ranking cluster, relevansi dihitung dari kombinasi share konten positif, share promotional, jumlah platform, dan jumlah entitas."
    )
    add_paragraph(
        doc,
        "Strategi split menggunakan holdout berbasis waktu: empat bulan terbaru menjadi data uji. Fitur graf yang dipakai pada klasifikasi dihitung secara historis sampai waktu sampel sehingga tidak melihat bukti masa depan. "
        "Snapshot ranking akhir tetap menggunakan seluruh data karena ia merepresentasikan kondisi saat sistem dijalankan untuk triase terkini."
    )
    source_small = sources[["source", "role", "status"]].copy()
    add_table(doc, source_small, "Tabel 3. Sumber resmi yang menjadi konteks dan rencana integrasi.", max_rows=6, widths=[2.1, 3.6, 0.9])

    add_heading(doc, "4. Metodologi", 1)
    add_paragraph(
        doc,
        "JudolFlow-X terdiri atas lima tahap. Tahap pertama melakukan normalisasi teks: lowercasing, penggantian karakter obfuscation umum, pembersihan URL, dan tokenisasi sederhana. "
        "Tahap kedua mengekstraksi entitas dengan regex konservatif untuk domain, handle, nomor kontak, dan petunjuk pembayaran. "
        "Tahap ketiga membentuk graf heterogen sample-entity dan entity-entity melalui ko-kemunculan dalam sampel yang sama."
    )
    add_paragraph(
        doc,
        "Tahap keempat melatih model StumpBoost, yaitu boosting ringan berbasis decision stump pada fitur teks, entitas, sumber, dan graf historis. "
        "Model ini dipilih karena mudah dijelaskan, tidak membutuhkan dependency berat, dan cukup kuat sebagai baseline kompetitif. "
        "Tahap kelima menggabungkan skor model, skor graf, jumlah entitas, jumlah domain, dan jangkauan platform menjadi priority_score cluster."
    )
    add_bullets(
        doc,
        [
            "Baseline acak mengukur performa tanpa informasi.",
            "Weighted index baseline memberi pembanding transparan berbasis bobot fitur.",
            "StumpBoost JudolFlow menjadi model utama yang belajar dari kombinasi fitur.",
            "Ablation study menghapus keluarga fitur untuk melihat kontribusi teks, entitas, graf, dan sumber.",
            "Bootstrap Top-K overlap menguji stabilitas ranking terhadap variasi sampel.",
        ],
    )
    add_figure(doc, FIGURE_DIR / "cluster_risk_quadrant.png", "Gambar 1. Kuadran risiko cluster berdasarkan sinyal teks dan bukti graf.", width=5.8)

    doc.add_page_break()
    add_heading(doc, "5. Eksperimen dan Hasil", 1)
    class_metrics = metrics[metrics["task"] == "content_classification"][
        ["model", "macro_f1", "pr_auc", "precision_at_k", "recall_at_k", "brier"]
    ]
    add_table(doc, class_metrics, "Tabel 4. Hasil klasifikasi konten pada holdout waktu.", widths=[1.8, 0.8, 0.8, 0.9, 0.9, 0.8])
    add_figure(doc, FIGURE_DIR / "model_comparison.png", "Gambar 2. Perbandingan PR-AUC model klasifikasi konten.", width=5.8)
    cluster_metrics = metrics[metrics["task"] == "cluster_ranking"][["metric", "value"]]
    add_table(doc, cluster_metrics, "Tabel 5. Hasil evaluasi ranking cluster.", widths=[2.8, 1.2])
    add_table(doc, ablation, "Tabel 6. Ablation study keluarga fitur.", widths=[1.9, 0.8, 0.8, 0.8])
    add_figure(doc, FIGURE_DIR / "feature_importance.png", "Gambar 3. Fitur teratas pada model utama.", width=5.8)

    add_heading(doc, "6. Analisis", 1)
    add_paragraph(
        doc,
        "Hasil eksperimen menunjukkan weighted index sudah kuat karena corpus prototipe memiliki sinyal tekstual dan entitas yang sengaja dibuat eksplisit untuk menguji alur end-to-end. "
        "Model utama memperbaiki kalibrasi dan konsistensi ranking, tetapi angka yang sangat tinggi tidak boleh dibaca sebagai klaim performa lapangan. "
        "Pada data nyata, variasi bahasa, noise OCR, duplikasi, dan konten edukasi yang mengutip istilah judi akan menurunkan performa sehingga validasi manusia tetap wajib."
    )
    add_paragraph(
        doc,
        "Feature importance memperlihatkan bahwa kombinasi sinyal teks dan entitas menjadi penentu utama. Fitur graf berperan pada ranking cluster karena mampu menaikkan prioritas cluster yang memiliki banyak bukti lintas sampel atau lintas platform. "
        "Ablation dipakai untuk memastikan model tidak bergantung pada satu fitur tunggal, sedangkan bootstrap Top-K overlap mengukur apakah cluster prioritas tetap muncul ketika sampel berubah."
    )
    add_table(doc, importance.head(8)[["feature", "importance", "feature_group"]], "Tabel 7. Delapan fitur terpenting.", widths=[2.4, 0.9, 1.0])
    add_table(doc, stability, "Tabel 8. Stabilitas ranking Top-20.", widths=[3.0, 1.0])
    top_rank = ranking[["rank", "cluster_id", "priority_score", "sample_count", "entity_count", "platform_count", "top_reasons"]].head(10)
    add_table(doc, top_rank, "Tabel 9. Sepuluh cluster prioritas tertinggi.", widths=[0.45, 0.75, 0.8, 0.75, 0.75, 0.75, 2.5])
    add_figure(doc, FIGURE_DIR / "top_priority_clusters.png", "Gambar 4. Skor prioritas sepuluh cluster teratas.", width=5.8)

    doc.add_page_break()
    add_heading(doc, "7. Studi Kasus dan Dashboard", 1)
    add_paragraph(
        doc,
        "Dashboard JudolFlow Watch dirancang untuk reviewer. Tampilan utama berisi Top-K review queue, profil cluster, alasan prioritas, contoh teks yang telah dimasking, distribusi platform, feature importance, ablation, dan fairness/subgroup check. "
        "Desain dashboard menegaskan batas etis: output adalah antrian tinjauan manusia, bukan vonis otomatis."
    )
    add_table(doc, cases[["rank", "cluster_id", "priority_score", "top_reasons"]], "Tabel 10. Studi kasus lima cluster prioritas.", widths=[0.55, 0.85, 0.95, 3.8])
    add_heading(doc, "Contoh Interpretasi", 2)
    for _, row in cases.head(3).iterrows():
        add_paragraph(doc, str(row["interpretation"]))
    add_figure(doc, FIGURE_DIR / "dashboard_preview.png", "Gambar 5. Pratinjau dashboard JudolFlow Watch.", width=6.1)

    add_heading(doc, "8. Keterbatasan dan Risiko", 1)
    add_bullets(
        doc,
        [
            "Corpus bawaan adalah data prototipe anonim, bukan bukti lapangan. Dataset final harus berasal dari sumber publik/legal dan dianotasi tim.",
            "Weak supervision rentan bias terhadap kata kunci. Konten edukasi yang mengutip istilah judi bisa menjadi false positive.",
            "Regex entity extraction belum cukup untuk OCR rusak, bahasa campuran, atau obfuscation baru.",
            "Ranking cluster dapat memprioritaskan jaringan besar; reviewer perlu melihat alasan fitur agar cluster kecil tetapi berbahaya tidak terlewat.",
            "Sistem tidak boleh mengungkap live illegal identifier karena dapat menimbulkan risiko penyebaran ulang atau membantu aktor menguji evasion.",
        ],
    )

    add_heading(doc, "9. Kesimpulan", 1)
    add_paragraph(
        doc,
        "JudolFlow-X mengubah masalah deteksi judi online dari klasifikasi tunggal menjadi ranking cluster berbasis multi-sinyal. "
        "Pipeline yang dibangun sudah mencakup dataset builder, ekstraksi entitas, graf ko-kemunculan, model klasifikasi, ranking cluster, ablation, stability check, case study, dashboard, serta technical report. "
        "Kekuatan utama pendekatan ini adalah kesesuaian dengan operasi tinjauan: reviewer dapat melihat cluster prioritas, alasan fitur, dan contoh yang telah dimasking."
    )
    add_paragraph(
        doc,
        "Pengembangan berikutnya adalah memasukkan data publik beranotasi, menambah OCR screenshot yang legal, menguji embedding retrieval, menerapkan hard-example mining, dan melakukan evaluasi antar-platform. "
        "Dengan validasi manusia yang kuat, JudolFlow-X dapat menjadi karya data mining yang relevan dengan kemandirian bangsa melalui perlindungan ruang digital Indonesia."
    )

    add_heading(doc, "Daftar Pustaka", 1)
    refs = [
        "Kementerian Komunikasi dan Digital Republik Indonesia. Dirjen Wasdig Minta Masyarakat Waspada Penipuan Permintaan Data Pribadi Terkait Judi Online. https://portal.komdigi.go.id/kanal-publik/berita-kini/9397",
        "Kementerian Komunikasi dan Digital Republik Indonesia. SAMAN Siap Operasi Penuh Oktober, 2,1 Juta Konten Judol Sudah Ditindak. https://portal.komdigi.go.id/kanal-publik/berita-kini/9623",
        "Kementerian Komunikasi dan Digital Republik Indonesia. 23 Ribu Rekening Judi Online Diblokir, Publik Diminta Aktif Melapor. https://portal.komdigi.go.id/kanal-publik/berita-kini/9712",
        "PPATK. Pemerintah Tegaskan Perang Total terhadap Judi Online dan Pencucian Uang. https://www.ppatk.go.id/news/read/1555/pemerintah-tegaskan-perang-total-terhadap-judi-online-dan-pencucian-uang.html",
        "Otoritas Jasa Keuangan. OJK Memerintahkan Bank untuk Memblokir Rekening yang Terlibat dalam Kegiatan Judi Online. https://www.ojk.go.id/id/berita-dan-kegiatan/siaran-pers/Documents/Pages/OJK-Memerintahkan-Bank-untuk-Memblokir-Rekening-yang-Terlibat-dalam-Kegiatan-Judi-Online/SP%20-%20OJK%20MEMERINTAHKAN%20BANK%20UNTUK%20MEMBLOKIR%20REKENING%20YANG%20TERLIBAT%20DALAM%20KEGIATAN%20JUDI%20ONLINE.pdf",
        "N. Reimers and I. Gurevych. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. EMNLP-IJCNLP, 2019. https://aclanthology.org/D19-1410/",
        "J. Devlin, M. Chang, K. Lee, and K. Toutanova. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. NAACL, 2019. https://aclanthology.org/N19-1423/",
        "W. L. Hamilton, R. Ying, and J. Leskovec. Inductive Representation Learning on Large Graphs. NeurIPS, 2017. https://papers.nips.cc/paper/6703-inductive-representation-learning-on-large-graphs",
    ]
    for idx, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.add_run(f"[{idx}] {ref}").font.size = Pt(9.5)

    doc.save(DOCX_PATH)


def export_pdf() -> bool:
    script = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{str(DOCX_PATH)}')
$doc.SaveAs([ref]'{str(PDF_PATH)}', [ref]17)
$doc.Close($false)
$word.Quit()
"""
    try:
        subprocess.run(["powershell", "-NoProfile", "-Command", script], check=True, capture_output=True, text=True, timeout=120)
        return PDF_PATH.exists()
    except Exception as exc:
        (REPORT_DIR / "pdf_export_error.log").write_text(f"Word COM export failed:\n{exc}\n", encoding="utf-8")
        return export_pdf_from_rendered_pages()


def find_render_docx_script() -> Path | None:
    base = Path.home() / ".codex" / "plugins" / "cache" / "openai-primary-runtime" / "documents"
    candidates = sorted(base.glob("*/skills/documents/render_docx.py"), reverse=True)
    return candidates[0] if candidates else None


def render_docx_pages() -> List[Path]:
    script = find_render_docx_script()
    if script is None:
        return []
    out_dir = REPORT_DIR / "rendered_docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            sys.executable,
            str(script),
            str(DOCX_PATH),
            "--output_dir",
            str(out_dir),
            "--renderer",
            "artifact-tool",
        ],
        check=True,
        capture_output=True,
        text=True,
        timeout=180,
    )

    def page_key(path: Path) -> int:
        match = re.search(r"page-(\d+)", path.stem)
        return int(match.group(1)) if match else 0

    return sorted(out_dir.glob("page-*.png"), key=page_key)


def export_pdf_from_rendered_pages() -> bool:
    try:
        pages = render_docx_pages()
        if not pages:
            return False
        images = []
        for page in pages:
            img = Image.open(page).convert("RGB")
            img.thumbnail((1240, 1754))
            canvas = Image.new("RGB", (1240, 1754), "white")
            x = (1240 - img.width) // 2
            y = (1754 - img.height) // 2
            canvas.paste(img, (x, y))
            images.append(canvas)
        images[0].save(PDF_PATH, "PDF", resolution=150.0, save_all=True, append_images=images[1:], quality=86)
        return PDF_PATH.exists()
    except Exception as exc:
        with (REPORT_DIR / "pdf_export_error.log").open("a", encoding="utf-8") as fh:
            fh.write(f"\nRendered-page PDF fallback failed:\n{exc}\n")
        return False


def main() -> None:
    write_report()
    ok = export_pdf()
    print(f"DOCX: {DOCX_PATH}")
    print(f"PDF: {PDF_PATH if ok else 'export failed; see reports/pdf_export_error.log'}")


if __name__ == "__main__":
    main()
