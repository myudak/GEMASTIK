from __future__ import annotations

import json
import math
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
TABLES = ROOT / "outputs" / "tables"
FIGURES = ROOT / "outputs" / "figures"
DATA = ROOT / "data"

DOCX_OUT = (
    REPORTS
    / "GEMASTIK XVIII Penambangan Data - ID_TIM - NAMA_TIM - PanganShock-X.docx"
)
PDF_OUT = (
    REPORTS
    / "GEMASTIK XVIII Penambangan Data - ID_TIM - NAMA_TIM - PanganShock-X.pdf"
)


def load_inputs() -> dict:
    return {
        "summary": json.loads((REPORTS / "dataset_summary.json").read_text(encoding="utf-8")),
        "metrics": json.loads((REPORTS / "final_metrics.json").read_text(encoding="utf-8")),
        "classification": pd.read_csv(TABLES / "classification_model_results.csv"),
        "regression": pd.read_csv(TABLES / "regression_model_results.csv"),
        "ablation": pd.read_csv(TABLES / "feature_ablation.csv"),
        "horizon": pd.read_csv(TABLES / "horizon_comparison.csv"),
        "threshold": pd.read_csv(TABLES / "threshold_comparison.csv"),
        "coverage": pd.read_csv(TABLES / "dataset_commodity_coverage.csv"),
        "importance": pd.read_csv(TABLES / "feature_importance.csv"),
        "alerts": pd.read_csv(TABLES / "latest_alert_ranking.csv"),
        "panel": pd.read_csv(DATA / "processed" / "panganshock_monthly_panel.csv"),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C2CC", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 16, "000000"),
        ("Heading 1", 13, "000000"),
        ("Heading 2", 11, "000000"),
        ("Heading 3", 10.5, "000000"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(4)
    styles["Heading 2"].paragraph_format.space_before = Pt(8)
    styles["Heading 2"].paragraph_format.space_after = Pt(2)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def p(doc: Document, text: str, style: str | None = None, align=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.6) if style is None else None
    paragraph.add_run(text)
    return paragraph


def bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def heading(doc: Document, number: str, text: str, level: int = 1) -> None:
    doc.add_paragraph(f"{number} {text}", style=f"Heading {level}")


def caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"


def add_table(
    doc: Document,
    df: pd.DataFrame,
    caption_text: str,
    widths: list[float] | None = None,
    max_rows: int | None = None,
    font_size: float = 8.7,
) -> None:
    if max_rows:
        df = df.head(max_rows).copy()
    caption(doc, caption_text)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, col in enumerate(df.columns):
        cell = header.cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(str(col))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = "Times New Roman"
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths:
            cell.width = Cm(widths[i])

    for ridx, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cell = cells[i]
            if ridx % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if df[col].dtype == object else WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(row[col]))
            run.font.size = Pt(font_size)
            run.font.name = "Times New Roman"
            if widths:
                cell.width = Cm(widths[i])
    doc.add_paragraph()


def add_figure(doc: Document, image_name: str, caption_text: str, width_cm: float = 13.5) -> None:
    path = FIGURES / image_name
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    caption(doc, caption_text)


def fmt_pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def fmt_num(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def fmt_rp(value: float) -> str:
    return f"Rp{value:,.0f}".replace(",", ".")


def prepare_tables(inputs: dict) -> dict[str, pd.DataFrame]:
    metrics = inputs["metrics"]
    classification = inputs["classification"]
    regression = inputs["regression"]
    ablation = inputs["ablation"]
    horizon = inputs["horizon"]
    threshold = inputs["threshold"]
    coverage = inputs["coverage"]
    importance = inputs["importance"]
    alerts = inputs["alerts"]
    panel = inputs["panel"]

    cls = classification[
        (classification["horizon_months"] == metrics["selected_horizon_months"])
        & (classification["threshold"] == metrics["selected_threshold"])
    ][["model", "pr_auc", "roc_auc", "recall_top10", "precision_top10", "brier", "cal_brier"]].copy()
    cls.columns = ["Model", "PR-AUC", "ROC-AUC", "Recall@10%", "Precision@10%", "Brier", "Brier Cal."]
    for col in cls.columns[1:]:
        cls[col] = cls[col].map(lambda x: fmt_num(float(x)))

    reg = regression[regression["horizon_months"] == metrics["selected_horizon_months"]][
        ["model", "mae", "rmse", "wmape"]
    ].copy()
    final_row = pd.DataFrame(
        [
            {
                "model": "HistGB Final Train+Val",
                "mae": metrics["mae"],
                "rmse": metrics["rmse"],
                "wmape": metrics["wmape"],
            }
        ]
    )
    reg = pd.concat([reg, final_row], ignore_index=True)
    reg.columns = ["Model", "MAE", "RMSE", "wMAPE"]
    reg["MAE"] = reg["MAE"].map(lambda x: f"{float(x):,.0f}")
    reg["RMSE"] = reg["RMSE"].map(lambda x: f"{float(x):,.0f}")
    reg["wMAPE"] = reg["wMAPE"].map(lambda x: fmt_num(float(x)))

    abl = ablation[["feature_set", "pr_auc", "roc_auc", "recall_top10", "precision_top10", "brier"]].copy()
    abl.columns = ["Feature Set", "PR-AUC", "ROC-AUC", "Recall@10%", "Precision@10%", "Brier"]
    for col in abl.columns[1:]:
        abl[col] = abl[col].map(lambda x: fmt_num(float(x)))
    abl["Feature Set"] = abl["Feature Set"].replace(
        {
            "price_only": "Harga + ID",
            "price_calendar": "Harga + Kalender",
            "price_calendar_weather": "Harga + Kalender + Cuaca",
            "all": "Semua fitur inti",
        }
    )

    hor = horizon.copy()
    hor.columns = ["Horizon (bulan)", "PR-AUC", "Recall@10%", "Brier", "MAE", "wMAPE"]
    for col in ["PR-AUC", "Recall@10%", "Brier", "wMAPE"]:
        hor[col] = hor[col].map(lambda x: fmt_num(float(x)))
    hor["MAE"] = hor["MAE"].map(lambda x: f"{float(x):,.0f}")

    thr = threshold.copy()
    thr.columns = ["Threshold", "Positive Rate", "PR-AUC", "Recall@10%", "Precision@10%", "Brier"]
    thr["Threshold"] = thr["Threshold"].map(lambda x: fmt_pct(float(x)))
    thr["Positive Rate"] = thr["Positive Rate"].map(lambda x: fmt_pct(float(x)))
    for col in ["PR-AUC", "Recall@10%", "Precision@10%", "Brier"]:
        thr[col] = thr[col].map(lambda x: fmt_num(float(x)))

    cov = coverage[["commodity", "rows", "non_missing", "missing", "coverage_rate", "mean_price"]].copy()
    cov = cov[cov["coverage_rate"] >= 0.70].sort_values("commodity")
    cov.columns = ["Komoditas", "Baris", "Non-missing", "Missing", "Coverage", "Rata-rata Harga"]
    cov["Coverage"] = cov["Coverage"].map(lambda x: fmt_pct(float(x)))
    cov["Rata-rata Harga"] = cov["Rata-rata Harga"].map(lambda x: fmt_rp(float(x)))

    imp = importance.head(12).copy()
    imp.columns = ["Fitur", "Importance"]
    imp["Importance"] = imp["Importance"].map(lambda x: fmt_num(float(x), 4))

    alert_cols = [
        "rank",
        "province",
        "commodity",
        "price_filled",
        "spike_probability_calibrated",
        "predicted_return",
        "priority_level",
        "top_driver_1",
    ]
    al = alerts[alert_cols].head(10).copy()
    al.columns = ["Rank", "Provinsi", "Komoditas", "Harga", "P(spike)", "Pred. Return", "Prioritas", "Driver Utama"]
    al["Harga"] = al["Harga"].map(lambda x: fmt_rp(float(x)))
    al["P(spike)"] = al["P(spike)"].map(lambda x: fmt_pct(float(x)))
    al["Pred. Return"] = al["Pred. Return"].map(lambda x: fmt_pct(float(x)))

    panel["date"] = pd.to_datetime(panel["date"])
    main_target = f"spike_h{int(metrics['selected_horizon_months'])}_t{int(metrics['selected_threshold'] * 100)}"
    commodity_error = (
        panel.dropna(subset=[main_target])
        .groupby("commodity")
        .agg(
            Observasi=(main_target, "size"),
            Spike_Rate=(main_target, "mean"),
            Harga_Rata2=("price_filled", "mean"),
        )
        .reset_index()
        .sort_values("Spike_Rate", ascending=False)
        .head(10)
    )
    commodity_error.columns = ["Komoditas", "Observasi", "Spike Rate", "Harga Rata-rata"]
    commodity_error["Spike Rate"] = commodity_error["Spike Rate"].map(lambda x: fmt_pct(float(x)))
    commodity_error["Harga Rata-rata"] = commodity_error["Harga Rata-rata"].map(lambda x: fmt_rp(float(x)))

    return {
        "classification": cls,
        "regression": reg,
        "ablation": abl,
        "horizon": hor,
        "threshold": thr,
        "coverage": cov,
        "importance": imp,
        "alerts": al,
        "commodity_error": commodity_error,
    }


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PanganShock-X: Prediksi Risiko Lonjakan Harga Pangan Strategis Berbasis Multiview Time-Series Data Mining"
    )
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(3)
    run = meta.add_run("Nama Tim: <NAMA TIM> | ID Tim: <ID TIM> | Perguruan Tinggi: <UNIVERSITAS>")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run("Anggota: <Anggota 1>, <Anggota 2>, <Anggota 3>")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    doc.add_paragraph()


def build_report() -> Path:
    inputs = load_inputs()
    tables = prepare_tables(inputs)
    summary = inputs["summary"]
    metrics = inputs["metrics"]

    doc = Document()
    style_document(doc)
    add_page_number(doc.sections[0])
    add_title_block(doc)

    heading(doc, "", "Abstrak", 1)
    p(
        doc,
        (
            "Lonjakan harga pangan strategis dapat mengganggu daya beli masyarakat dan menghambat "
            "respons intervensi pasar apabila baru diketahui setelah terjadi. Penelitian ini mengusulkan "
            "PanganShock-X, pipeline penambangan data untuk memprediksi risiko lonjakan harga pangan "
            "pada level provinsi-komoditas dan mengubah prediksi tersebut menjadi ranking prioritas. "
            f"Dataset utama berasal dari Portal Data Badan Pangan Nasional periode {summary['start_date']} "
            f"sampai {summary['end_date']} dengan {summary['provinces']} provinsi dan {summary['commodities']} "
            "komoditas strategis, kemudian diperkaya fitur kalender dan cuaca historis. Karena data terbuka "
            "yang konsisten tersedia pada granularitas bulanan, implementasi ini menggunakan horizon 1 dan "
            "2 bulan dengan target utama lonjakan minimal 5% dalam 2 bulan. Evaluasi dilakukan menggunakan "
            "time-based split, PR-AUC, Recall@Top-K, Brier score, MAE, RMSE, dan wMAPE. Model final mencapai "
            f"PR-AUC {metrics['pr_auc']:.3f}, ROC-AUC {metrics['roc_auc']:.3f}, Recall@Top-10% "
            f"{metrics['recall_top10']:.3f}, Precision@Top-10% {metrics['precision_top10']:.3f}, "
            f"Brier terkalibrasi {metrics['cal_brier']:.3f}, MAE {metrics['mae']:.0f}, dan wMAPE "
            f"{metrics['wmape']:.3f}. Output utama berupa daftar prioritas provinsi-komoditas dengan "
            "probabilitas lonjakan, estimasi kenaikan harga, skor prioritas, dan faktor penjelas."
        ),
    )
    kw = doc.add_paragraph()
    kw.alignment = WD_ALIGN_PARAGRAPH.LEFT
    kw.add_run("Kata kunci: ").bold = True
    kw.add_run("harga pangan, time-series, spike detection, early warning, calibrated ranking, data mining")

    heading(doc, "1.", "Pendahuluan", 1)
    heading(doc, "1.1", "Latar Belakang", 2)
    p(
        doc,
        (
            "Kemandirian bangsa pada sektor TIK tidak hanya berkaitan dengan pembangunan perangkat lunak, "
            "tetapi juga kemampuan memanfaatkan data untuk mendukung keputusan publik. Harga pangan strategis "
            "seperti beras, cabai, bawang, telur, daging ayam, gula, dan minyak goreng merupakan indikator "
            "yang dekat dengan kesejahteraan rumah tangga. Kenaikan yang tajam dapat menekan daya beli, "
            "memengaruhi inflasi daerah, dan menimbulkan kebutuhan intervensi seperti monitoring stok, "
            "penguatan distribusi, atau operasi pasar."
        ),
    )
    p(
        doc,
        (
            "Pemantauan harga pangan umumnya menampilkan harga terkini atau tren historis. Pendekatan tersebut "
            "berguna untuk observasi, tetapi belum langsung menjawab pertanyaan operasional: provinsi-komoditas "
            "mana yang perlu diprioritaskan karena berisiko mengalami lonjakan dalam waktu dekat? Oleh karena itu, "
            "PanganShock-X memformulasikan masalah sebagai kombinasi prediksi perubahan harga, klasifikasi "
            "lonjakan, dan ranking prioritas."
        ),
    )
    heading(doc, "1.2", "Rumusan Masalah dan Tujuan", 2)
    bullet(doc, "Bagaimana memprediksi risiko lonjakan harga pangan strategis pada level provinsi-komoditas?")
    bullet(doc, "Bagaimana membangun target lonjakan yang jelas dan bebas leakage pada data time-series?")
    bullet(doc, "Model dan fitur apa yang paling berkontribusi terhadap performa early warning?")
    bullet(doc, "Bagaimana mengubah probabilitas model menjadi ranking prioritas yang mudah dijelaskan?")
    p(
        doc,
        (
            "Tujuan penelitian adalah membangun pipeline yang reproducible, mengevaluasi baseline dan model utama, "
            "menyajikan ablation study, serta menghasilkan tabel prioritas yang dapat digunakan sebagai dasar "
            "monitoring. Manfaatnya adalah memperkuat analitik pangan, mendukung prioritas observasi, dan memberi "
            "contoh pemanfaatan data mining untuk kemandirian pangan."
        ),
    )
    heading(doc, "1.3", "Batasan", 2)
    p(
        doc,
        (
            "Batasan utama terletak pada granularitas data terbuka. Dataset Bapanas yang dapat direproduksi secara "
            "konsisten tersedia pada level bulanan, sehingga horizon implementasi adalah 1-2 bulan, bukan 7-14 hari. "
            "Data cuaca historis menggunakan Open-Meteo sebagai fallback terbuka berdasarkan koordinat ibu kota "
            "provinsi. Versi operasional dapat ditingkatkan menggunakan data harian PIHPS/BI dan data historis BMKG "
            "apabila akses resmi tersedia."
        ),
    )

    heading(doc, "2.", "Kajian Terkait", 1)
    p(
        doc,
        (
            "Forecasting harga komoditas pangan telah banyak menggunakan model time-series dan machine learning. "
            "Namun, prediksi nilai harga tidak selalu cukup untuk kebutuhan intervensi karena pengambil keputusan "
            "lebih membutuhkan daftar prioritas risiko. Oleh sebab itu, pendekatan ini memadukan forecasting, "
            "imbalanced classification, dan ranking."
        ),
    )
    p(
        doc,
        (
            "Random Forest [5] dan gradient boosting [6], [7] banyak digunakan pada data tabular karena kuat "
            "terhadap hubungan non-linear dan interaksi fitur. Untuk kasus kelas tidak seimbang, PR-AUC dan "
            "precision-recall curve lebih informatif dibanding akurasi atau ROC saja [8]. Karena output dipakai "
            "sebagai skor risiko, kalibrasi probabilitas juga penting agar skor lebih dapat ditafsirkan [9]. "
            "Perbedaan utama PanganShock-X adalah penekanan pada spike-risk ranking, bukan hanya regresi harga."
        ),
    )

    heading(doc, "3.", "Dataset dan Formulasi Masalah", 1)
    heading(doc, "3.1", "Sumber Data", 2)
    p(
        doc,
        (
            "Dataset harga berasal dari Portal Data Badan Pangan Nasional, yaitu rata-rata harga pangan bulanan "
            "tingkat konsumen provinsi. Dataset ini dipilih karena terbuka, resmi, dan memiliki cakupan nasional. "
            "Data cuaca historis diperoleh dari Open-Meteo Historical Weather API berdasarkan koordinat ibu kota "
            "provinsi. Fitur kalender dibuat secara deterministik dari bulan, tahun, dan penanda jendela Ramadan/"
            "Idul Fitri serta akhir tahun."
        ),
    )
    ds = pd.DataFrame(
        [
            ["Baris raw", f"{summary['raw_rows']:,}"],
            ["Baris panel terproses", f"{summary['processed_rows']:,}"],
            ["Rentang waktu", f"{summary['start_date']} sampai {summary['end_date']}"],
            ["Provinsi", str(summary["provinces"])],
            ["Komoditas strategis", str(summary["commodities"])],
            ["Missing harga raw", f"{summary['missing_price_raw_pct']:.2f}%"],
            ["Unit analisis", "provinsi x komoditas x bulan"],
        ],
        columns=["Komponen", "Nilai"],
    )
    add_table(doc, ds, "Tabel 1. Ringkasan dataset PanganShock-X.", widths=[6, 9], font_size=9)
    add_table(doc, tables["coverage"], "Tabel 2. Komoditas strategis yang digunakan dalam eksperimen.", font_size=8.2)

    heading(doc, "3.2", "Definisi Target", 2)
    p(
        doc,
        (
            "Untuk setiap baris pada waktu t, target regresi adalah return harga pada t+h. Target klasifikasi "
            "dibangun dari maksimum harga pada horizon ke depan. Sebuah observasi diberi label lonjakan jika "
            "harga maksimum antara t+1 sampai t+h naik minimal threshold dibanding harga pada t. Threshold 5%, "
            "8%, dan 10% diuji; konfigurasi utama menggunakan 5% karena menghasilkan jumlah kasus positif yang "
            "cukup untuk evaluasi dan tetap mencerminkan perubahan harga yang bermakna pada data bulanan."
        ),
    )
    formulas = pd.DataFrame(
        [
            ["Regresi", "return_h = (price(t+h) - price(t)) / price(t)"],
            ["Klasifikasi", "spike_h = 1 jika max(price(t+1..t+h)) >= price(t) x (1 + threshold)"],
            ["Ranking", "priority_score = P(spike) terkalibrasi x max(predicted_return, 0)"],
        ],
        columns=["Komponen", "Definisi"],
    )
    add_table(doc, formulas, "Tabel 3. Definisi target dan skor prioritas.", widths=[4, 11], font_size=8.8)

    heading(doc, "3.3", "Split Data Berbasis Waktu", 2)
    p(
        doc,
        (
            "Untuk mencegah leakage, data tidak dibagi secara acak. Train menggunakan periode 2021-2023, "
            "validation menggunakan 2024, dan test menggunakan observasi 2025 yang masih memiliki label horizon. "
            "Semua fitur rolling dan lag hanya menggunakan informasi sampai waktu t. Fitur kalender diperbolehkan "
            "karena tanggal hari besar diketahui sebelum periode prediksi."
        ),
    )
    split = pd.DataFrame(
        [
            ["Train", "2021-01 s.d. 2023-12", f"{metrics['n_train']:,}"],
            ["Validation", "2024-01 s.d. 2024-12", f"{metrics['n_validation']:,}"],
            ["Test", "2025-01 s.d. 2025-04", f"{metrics['n_test']:,}"],
        ],
        columns=["Split", "Periode", "Jumlah Observasi"],
    )
    add_table(doc, split, "Tabel 4. Pembagian data berbasis waktu untuk konfigurasi utama.", widths=[3.5, 6, 4], font_size=9)

    heading(doc, "4.", "Metodologi", 1)
    heading(doc, "4.1", "Preprocessing dan Feature Engineering", 2)
    p(
        doc,
        (
            "Preprocessing meliputi standardisasi nama provinsi dan komoditas, parsing nilai rupiah, penghapusan "
            "duplikasi, pengisian gap pendek dengan forward/backward fill terbatas, dan penggabungan fitur cuaca. "
            "Fitur utama meliputi lag harga 1, 2, 3, 6, dan 12 bulan; return historis; rolling mean, rolling "
            "standard deviation, rolling min/max; momentum; gap harga provinsi terhadap rata-rata nasional "
            "komoditas; peringkat harga antarprovinsi; fitur bulan siklik; dan indikator periode khusus."
        ),
    )
    heading(doc, "4.2", "Model", 2)
    p(
        doc,
        (
            "Baseline regresi mencakup persistence, moving average 3 bulan, dan Ridge Regression. Baseline "
            "klasifikasi menggunakan Logistic Regression dengan class weight. Model utama menggunakan Random "
            "Forest dan HistGradientBoosting. Probabilitas dikalibrasi pada validation set menggunakan Platt "
            "scaling. Model final dipilih berdasarkan kombinasi PR-AUC, Recall@Top-10%, precision alert, Brier "
            "score, dan interpretabilitas hasil ranking."
        ),
    )
    heading(doc, "4.3", "Metrik Evaluasi", 2)
    p(
        doc,
        (
            "PR-AUC digunakan sebagai metrik utama klasifikasi karena kasus lonjakan relatif tidak seimbang. "
            "Recall@Top-10% mengukur seberapa banyak lonjakan aktual yang tertangkap oleh 10% alert teratas, "
            "sedangkan Precision@Top-10% mengukur kualitas alert prioritas. Brier score mengevaluasi kalibrasi "
            "probabilitas. Untuk regresi digunakan MAE, RMSE, dan wMAPE agar kesalahan harga dapat dibaca dalam "
            "satuan rupiah maupun proporsi."
        ),
    )

    heading(doc, "5.", "Hasil Eksperimen", 1)
    add_table(doc, tables["horizon"], "Tabel 5. Perbandingan horizon prediksi.", font_size=8.5)
    add_table(doc, tables["threshold"], "Tabel 6. Perbandingan threshold label lonjakan.", font_size=8.5)
    add_table(doc, tables["classification"], "Tabel 7. Hasil klasifikasi risiko lonjakan pada konfigurasi utama.", font_size=8.3)
    add_table(doc, tables["regression"], "Tabel 8. Hasil regresi perubahan harga pada horizon utama.", font_size=8.2)
    add_table(doc, tables["ablation"], "Tabel 9. Ablation study kelompok fitur.", font_size=8.2)
    add_figure(
        doc,
        "classification_model_comparison.png",
        "Gambar 1. Model non-linear menghasilkan kualitas alert yang lebih tinggi daripada Logistic Regression.",
        width_cm=13.0,
    )
    add_figure(
        doc,
        "precision_recall_curve.png",
        "Gambar 2. Precision-recall curve model final untuk tugas deteksi lonjakan.",
        width_cm=11.5,
    )
    add_figure(
        doc,
        "calibration_curve.png",
        "Gambar 3. Kalibrasi probabilitas membantu menurunkan Brier score walau ranking tetap sama.",
        width_cm=11.5,
    )

    heading(doc, "6.", "Analisis Hasil", 1)
    heading(doc, "6.1", "Analisis Model", 2)
    p(
        doc,
        (
            f"Pada konfigurasi utama, positive rate test adalah {fmt_pct(metrics['test_positive_rate'])}. "
            f"HistGradientBoosting mencapai PR-AUC {metrics['pr_auc']:.3f} dan ROC-AUC {metrics['roc_auc']:.3f}. "
            f"Recall@Top-10% sebesar {metrics['recall_top10']:.3f} berarti 10% alert teratas menangkap sekitar "
            f"{metrics['recall_top10'] * 100:.1f}% kejadian lonjakan aktual pada periode uji. Precision@Top-10% "
            f"sebesar {metrics['precision_top10']:.3f} menunjukkan bahwa lebih dari separuh alert prioritas "
            "merupakan kasus positif. Ini lebih relevan untuk skenario monitoring terbatas dibanding sekadar akurasi."
        ),
    )
    p(
        doc,
        (
            "Ablation study menunjukkan bahwa fitur kalender dan cuaca menambah sinyal di atas fitur harga saja. "
            "Price-only sudah cukup kuat karena harga pangan memiliki autokorelasi dan momentum, tetapi penambahan "
            "kalender dan cuaca menaikkan PR-AUC serta Recall@Top-10%. Hal ini mendukung klaim bahwa pendekatan "
            "multiview lebih sesuai daripada hanya memakai satu ruang fitur."
        ),
    )
    heading(doc, "6.2", "Feature Importance dan Error Analysis", 2)
    add_table(doc, tables["importance"], "Tabel 10. Dua belas fitur terpenting menurut Random Forest.", font_size=8.2)
    add_table(doc, tables["commodity_error"], "Tabel 11. Komoditas dengan spike rate tertinggi pada panel utama.", font_size=8.2)
    add_figure(
        doc,
        "feature_importance.png",
        "Gambar 4. Feature importance memperlihatkan dominasi sinyal harga historis, momentum, dan fitur wilayah/komoditas.",
        width_cm=13.2,
    )
    heading(doc, "6.3", "Analisis Ranking Prioritas", 2)
    p(
        doc,
        (
            "Ranking prioritas tidak hanya menggunakan probabilitas lonjakan. Skor dihitung dari probabilitas "
            "terkalibrasi dikalikan prediksi return positif, sehingga alert teratas merepresentasikan kombinasi "
            "risiko dan potensi dampak harga. Kolom driver utama dibuat dari fitur yang menonjol, seperti volatilitas "
            "terbaru, momentum harga, harga relatif terhadap rata-rata nasional, dan jendela Ramadan/Idul Fitri."
        ),
    )
    doc.add_page_break()
    add_table(doc, tables["alerts"], "Tabel 12. Contoh ranking prioritas provinsi-komoditas terbaru.", font_size=7.6)
    add_figure(
        doc,
        "alert_heatmap.png",
        "Gambar 5. Heatmap alert prioritas membantu melihat konsentrasi risiko pada komoditas dan wilayah tertentu.",
        width_cm=13.2,
    )
    heading(doc, "6.4", "Keterbatasan", 2)
    bullet(doc, "Granularitas data terbuka adalah bulanan, sehingga sistem belum menjadi alert harian 7-14 hari.")
    bullet(doc, "Fitur cuaca menggunakan Open-Meteo sebagai fallback; versi final sebaiknya memakai data historis BMKG bila tersedia.")
    bullet(doc, "Model belum memasukkan stok, distribusi, produksi, atau berita karena akses data yang konsisten belum dipastikan.")
    bullet(doc, "Analisis perlu diverifikasi oleh tim terhadap peristiwa harga pangan aktual pada periode uji.")

    heading(doc, "7.", "Kesimpulan dan Pengembangan Lanjutan", 1)
    p(
        doc,
        (
            "PanganShock-X berhasil diimplementasikan sebagai pipeline penambangan data end-to-end untuk prediksi "
            "risiko lonjakan harga pangan strategis dan penyusunan ranking prioritas provinsi-komoditas. Kontribusi "
            "utama penelitian adalah formulasi spike-risk ranking, dataset panel multiview, validasi berbasis waktu, "
            "evaluasi dengan metrik alert yang relevan, serta output yang dapat dijelaskan. Pengembangan lanjutan "
            "meliputi integrasi data harian PIHPS/BI, data BMKG historis, fitur stok/distribusi, serta rolling-origin "
            "backtesting yang lebih luas."
        ),
    )

    heading(doc, "8.", "Daftar Pustaka", 1)
    refs = [
        "[1] Panitia GEMASTIK XVIII, Panduan GEMASTIK 2025, 2025.",
        "[2] Badan Pangan Nasional, Rata-rata Harga Pangan Bulanan Tingkat Konsumen Provinsi (Angka Juni 2025), Portal Data Badan Pangan Nasional, 2025.",
        "[3] Open-Meteo, Historical Weather API Documentation.",
        "[4] R. J. Hyndman and G. Athanasopoulos, Forecasting: Principles and Practice, 3rd ed., OTexts, 2021.",
        "[5] L. Breiman, Random Forests, Machine Learning, vol. 45, pp. 5-32, 2001.",
        "[6] J. H. Friedman, Greedy Function Approximation: A Gradient Boosting Machine, Annals of Statistics, vol. 29, no. 5, pp. 1189-1232, 2001.",
        "[7] G. Ke et al., LightGBM: A Highly Efficient Gradient Boosting Decision Tree, NeurIPS, 2017.",
        "[8] J. Davis and M. Goadrich, The Relationship Between Precision-Recall and ROC Curves, ICML, 2006.",
        "[9] A. Niculescu-Mizil and R. Caruana, Predicting Good Probabilities with Supervised Learning, ICML, 2005.",
        "[10] F. Pedregosa et al., Scikit-learn: Machine Learning in Python, Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
    ]
    for ref in refs:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.6)
        para.paragraph_format.first_line_indent = Cm(-0.6)
        para.paragraph_format.space_after = Pt(2)
        para.add_run(ref)

    heading(doc, "Lampiran A.", "Reproducibility Ringkas", 1)
    p(
        doc,
        (
            "Pipeline dapat dijalankan ulang dari folder PanganShock-X dengan membuat virtual environment, memasang "
            "requirements.txt, lalu menjalankan `python src/panganshock_pipeline.py`. Script akan mengunduh/cache "
            "data Bapanas dan Open-Meteo, membangun panel, melatih model, menyimpan tabel, membuat figur, dan "
            "menghasilkan draft laporan."
        ),
    )

    doc.save(DOCX_OUT)
    return DOCX_OUT


def export_pdf_with_word(docx_path: Path, pdf_path: Path) -> None:
    ps = f"""
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{str(docx_path)}')
$doc.ExportAsFixedFormat('{str(pdf_path)}', 17)
$doc.Close($false)
$word.Quit()
"""
    subprocess.run(["powershell", "-NoProfile", "-Command", ps], check=True)


def main() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)
    docx_path = build_report()
    export_pdf_with_word(docx_path, PDF_OUT)
    print(f"DOCX: {docx_path}")
    print(f"PDF: {PDF_OUT}")


if __name__ == "__main__":
    main()
