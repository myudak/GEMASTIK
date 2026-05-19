from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
PROJECTS_ROOT = ROOT.parent
TEMPLATE = PROJECTS_ROOT / "[Template] Makalah Gemastik ieee.docx"
REPORT_DIR = ROOT / "reports"
TABLE_DIR = REPORT_DIR / "tables"
FIGURE_DIR = REPORT_DIR / "figures"
DOCX_PATH = REPORT_DIR / "GEMASTIK XVIII Penambangan Data - ID_TIM - NAMA_TIM - PanganFlow.ID.docx"
PDF_PATH = REPORT_DIR / "GEMASTIK XVIII Penambangan Data - ID_TIM - NAMA_TIM - PanganFlow.ID.pdf"


def load_inputs() -> dict[str, object]:
    required = {
        "dataset_summary": TABLE_DIR / "dataset_summary.csv",
        "metrics": TABLE_DIR / "metrics_summary.csv",
        "ranking": TABLE_DIR / "ranking_snapshot.csv",
        "flows": TABLE_DIR / "flow_recommendations.csv",
        "importance": TABLE_DIR / "feature_importance.csv",
        "ablation": TABLE_DIR / "feature_ablation.csv",
        "source_status": TABLE_DIR / "source_status.csv",
        "validation": TABLE_DIR / "validation_checks.csv",
        "clusters": TABLE_DIR / "cluster_summary.csv",
        "fairness": TABLE_DIR / "subgroup_fairness.csv",
        "stability": TABLE_DIR / "ranking_stability.csv",
        "cases": TABLE_DIR / "case_studies.csv",
    }
    missing = [str(path) for path in required.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("Run build_dataset.py and train_panganflow.py first. Missing: " + ", ".join(missing))
    return {key: pd.read_csv(path) for key, path in required.items()}


def clear_document(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)
    set_columns(section, 1, Cm(0.65))

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    for name in ["IEEEParagraph", "IEEEAbtract", "IEEEReferenceItem"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
            styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            styles[name].font.size = Pt(9)
            styles[name].paragraph_format.space_after = Pt(2)

    for name in ["GemastikHeading1", "IEEEHeading1", "Heading 1"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
            styles[name].font.bold = True
            styles[name].font.size = Pt(9.4)
            styles[name].font.color.rgb = RGBColor(0, 0, 0)
            styles[name].paragraph_format.space_before = Pt(8)
            styles[name].paragraph_format.space_after = Pt(3)

    for name in ["GemastikHeading2", "IEEEHeading2", "Heading 2"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
            styles[name].font.bold = True
            styles[name].font.size = Pt(9.2)
            styles[name].font.color.rgb = RGBColor(0, 0, 0)
            styles[name].paragraph_format.space_before = Pt(5)
            styles[name].paragraph_format.space_after = Pt(2)


def set_columns(section, count: int, space: Cm) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(int(space.cm * 567)))


def add_page_numbers(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_name(doc: Document, preferred: Iterable[str], fallback: str) -> str:
    available = {style.name for style in doc.styles}
    for name in preferred:
        if name in available:
            return name
    return fallback if fallback in available else "Normal"


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PanganFlow.ID: Penambangan Pola Surplus-Defisit dan Rekomendasi Redistribusi Beras Antarprovinsi")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(17)

    authors = [
        "Nama Ketua Tim1, Nama Anggota Tim 12, Nama Anggota Tim 23, Nama Pembimbing4",
        "1,2,3Program Studi / Fakultas, Nama Universitas, Kota, Indonesia",
        "4Program Studi / Fakultas, Nama Universitas, Kota, Indonesia",
        "Corresponding Author: Nama Ketua Tim, email: ketuatim@institusi.ac.id",
    ]
    for text in authors:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(8.5)


def add_abstract(doc: Document, summary: pd.DataFrame) -> None:
    row = summary.iloc[0]
    text = (
        "INTISARI - Ketahanan pangan tidak hanya ditentukan oleh kecukupan stok nasional, tetapi juga oleh ketimpangan harga, produksi, kebutuhan, dan akses antarwilayah. "
        "Penelitian ini mengusulkan PanganFlow.ID, sistem penambangan data untuk mendeteksi provinsi dengan tekanan pasokan beras dan menyusun prioritas aliran redistribusi antarprovinsi. "
        f"Dataset yang dibangun mencakup {int(row['rows']):,} baris panel provinsi-bulan untuk {int(row['province_count'])} provinsi pada periode {row['date_min']} sampai {row['date_max']}. "
        "Fitur yang digunakan meliputi harga beras medium, gap terhadap median nasional, perubahan harga, rolling anomaly score, produksi beras, estimasi kebutuhan, proxy surplus-defisit, dan jarak logistik antarprovinsi. "
        "Metode yang diuji meliputi baseline acak, baseline harga, baseline defisit, weighted priority index, model Random Forest, K-Means clustering, Isolation Forest, dan graph-based redistribution scoring. "
        "Hasil sistem berupa ranking provinsi prioritas, tipologi wilayah, serta rekomendasi pasangan provinsi surplus menuju provinsi defisit yang dapat dipakai sebagai bahan tinjauan kebijakan."
    )
    p = doc.add_paragraph(style=style_name(doc, ["IEEEAbtract", "IEEE Paragraph"], "Normal"))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("KATA KUNCI - Data Mining, Beras, Ketahanan Pangan, Redistribusi, Anomaly Detection, Ranking.")
    r.bold = True
    r.font.size = Pt(9)


def begin_ieee_body(doc: Document) -> None:
    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    body_section.page_width = Cm(21.0)
    body_section.page_height = Cm(29.7)
    body_section.top_margin = Cm(2.3)
    body_section.bottom_margin = Cm(1.8)
    body_section.left_margin = Cm(1.3)
    body_section.right_margin = Cm(1.3)
    set_columns(body_section, 2, Cm(0.42))


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    if level == 1:
        style = style_name(doc, ["GemastikHeading1", "IEEEHeading1"], "Heading 1")
    else:
        style = style_name(doc, ["GemastikHeading2", "IEEEHeading2"], "Heading 2")
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(doc: Document, text: str) -> None:
    style = style_name(doc, ["IEEEParagraph"], "Normal")
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.35)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=style_name(doc, ["List Bullet", "IEEEBullet1"], "Normal"))
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(item if p.style.name != "Normal" else f"- {item}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(8.8)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(7.8)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=60, bottom=60, end=60) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C2CC", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_table(doc: Document, df: pd.DataFrame, caption_text: str, max_rows: int | None = None, font_size: float = 6.8) -> None:
    if max_rows:
        df = df.head(max_rows).copy()
    font_size = min(font_size, 7.2)
    caption(doc, caption_text)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 8.25)
    set_table_borders(table)
    col_width = Cm(8.25 / max(len(df.columns), 1))
    header = table.rows[0]
    for i, col in enumerate(df.columns):
        cell = header.cells[i]
        cell.width = col_width
        set_cell_shading(cell, "E7EEF2")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(col))
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(font_size)
    for ridx, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cell = cells[i]
            cell.width = col_width
            if ridx % 2 == 1:
                set_cell_shading(cell, "F8FAFB")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if df[col].dtype == object else WD_ALIGN_PARAGRAPH.CENTER
            value = row[col]
            if isinstance(value, float):
                text = f"{value:.3f}" if abs(value) < 1000 else f"{value:,.0f}"
            else:
                text = str(value)
            r = p.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption_text: str, width: float = 3.15) -> None:
    if not path.exists():
        return
    width = min(width, 3.25)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, caption_text)


def simple_metric(metrics: pd.DataFrame, model: str, metric: str) -> float:
    return float(metrics.loc[metrics["model"].eq(model), metric].iloc[0])


def compact_sources(df: pd.DataFrame) -> pd.DataFrame:
    cols = ["source", "status", "rows"]
    out = df[cols].copy()
    out["source"] = out["source"].astype(str).str.replace(" Monthly Consumer Price", "", regex=False)
    out["source"] = out["source"].astype(str).str.replace("PIHPS ", "", regex=False)
    return out


def write_report() -> None:
    data = load_inputs()
    summary: pd.DataFrame = data["dataset_summary"]  # type: ignore[assignment]
    metrics: pd.DataFrame = data["metrics"]  # type: ignore[assignment]
    ranking: pd.DataFrame = data["ranking"]  # type: ignore[assignment]
    flows: pd.DataFrame = data["flows"]  # type: ignore[assignment]
    importance: pd.DataFrame = data["importance"]  # type: ignore[assignment]
    ablation: pd.DataFrame = data["ablation"]  # type: ignore[assignment]
    source_status: pd.DataFrame = data["source_status"]  # type: ignore[assignment]
    validation: pd.DataFrame = data["validation"]  # type: ignore[assignment]
    clusters: pd.DataFrame = data["clusters"]  # type: ignore[assignment]
    fairness: pd.DataFrame = data["fairness"]  # type: ignore[assignment]
    stability: pd.DataFrame = data["stability"]  # type: ignore[assignment]
    cases: pd.DataFrame = data["cases"]  # type: ignore[assignment]

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE)) if TEMPLATE.exists() else Document()
    clear_document(doc)
    style_document(doc)
    add_page_numbers(doc.sections[0])

    add_title(doc)
    add_abstract(doc, summary)
    begin_ieee_body(doc)

    add_heading(doc, "I. PENDAHULUAN")
    add_paragraph(
        doc,
        "Indonesia sering dipahami memiliki kecukupan pangan pada level nasional, namun disparitas harga menunjukkan bahwa masalah pangan tidak berhenti pada total produksi. "
        "Harga beras dapat stabil di provinsi sentra produksi, sementara provinsi lain menghadapi harga lebih tinggi karena defisit produksi lokal, jarak distribusi, pola konsumsi, dan waktu panen. "
        "Kondisi kepulauan Indonesia membuat ketimpangan pasokan menjadi masalah spasial sekaligus temporal."
    )
    add_paragraph(
        doc,
        "Tujuan penelitian ini adalah membangun sistem data mining yang menjawab tiga pertanyaan: provinsi mana yang menunjukkan tekanan pasokan beras, provinsi mana yang relatif surplus, dan pasangan asal-tujuan mana yang layak ditinjau lebih dulu untuk redistribusi. "
        "Manfaat praktisnya adalah membantu penyusunan prioritas analisis, bukan menggantikan keputusan logistik pemerintah."
    )

    add_heading(doc, "II. KAJIAN TERKAIT")
    add_paragraph(
        doc,
        "PIHPS BI dikembangkan untuk mengurangi kesenjangan informasi harga, menjadi sinyal awal gejolak harga, dan membantu memahami keterkaitan harga antarprovinsi. "
        "Dalam literatur data mining, masalah sejenis umumnya didekati melalui deteksi anomali time series, clustering wilayah, forecasting harga, dan scoring prioritas. "
        "Kebaruan PanganFlow.ID adalah menggabungkan sinyal harga dengan proxy neraca produksi-kebutuhan dan scoring graf asal-tujuan, sehingga outputnya bukan hanya prediksi harga tetapi rekomendasi prioritas aliran pangan."
    )

    add_heading(doc, "III. DATASET DAN FORMULASI MASALAH")
    row = summary.iloc[0]
    add_paragraph(
        doc,
        f"Dataset final berbentuk panel provinsi-bulan-komoditas sebanyak {int(row['rows']):,} baris, mencakup {int(row['province_count'])} provinsi dan periode {row['date_min']} sampai {row['date_max']}. "
        "Komoditas v1 dibatasi pada Beras Medium agar perhitungan surplus-defisit konsisten dengan data produksi beras BPS dan konsumsi beras per kapita."
    )
    add_table(doc, compact_sources(source_status), "Tabel I. Status sumber data resmi dan snapshot.", max_rows=6, font_size=6.7)
    add_paragraph(
        doc,
        "Kebutuhan beras diproksikan sebagai populasi dikalikan konsumsi beras per kapita tahunan. Surplus proxy dihitung sebagai produksi beras tahunan dikurangi kebutuhan tahunan. "
        "Karena stok awal, arus masuk, arus keluar, dan cadangan tidak selalu tersedia terbuka pada resolusi provinsi-bulan, nilai ini diperlakukan sebagai proxy struktural, bukan neraca resmi final."
    )

    add_heading(doc, "IV. METODOLOGI")
    add_paragraph(
        doc,
        "Pipeline terdiri atas enam tahap. Pertama, harga beras diproses menjadi fitur gap terhadap median nasional, perubahan bulanan, rolling mean, rolling standard deviation, dan z-score historis. "
        "Kedua, produksi dan kebutuhan digabung menjadi skor defisit dan skor surplus. Ketiga, Isolation Forest digunakan untuk sinyal anomali harga. "
        "Keempat, Random Forest mempelajari prioritas pemantauan bulan berikutnya dari fitur bulan berjalan. Kelima, K-Means dan PCA digunakan untuk memetakan tipologi provinsi. "
        "Keenam, setiap pasangan provinsi surplus-defisit diberi skor graf berbasis price gap, surplus asal, defisit tujuan, skor model tujuan, anomali tujuan, dan penalti jarak."
    )
    add_bullets(
        doc,
        [
            "Baseline 1: ranking acak untuk mengukur batas bawah.",
            "Baseline 2: price-gap ranking untuk melihat efek disparitas harga saja.",
            "Baseline 3: deficit ranking untuk melihat efek neraca saja.",
            "Baseline 4: weighted priority index transparan sebagai pembanding model.",
            "Model utama: Random Forest ranker dengan split berbasis waktu.",
        ],
    )

    add_heading(doc, "V. EKSPERIMEN DAN HASIL")
    metric_view = metrics[["model", "ndcg_at_k", "precision_at_k", "recall_at_k"]].copy()
    metric_view["model"] = metric_view["model"].str.replace("Weighted priority index", "Weighted index", regex=False)
    metric_view["model"] = metric_view["model"].str.replace("PanganFlow hybrid model", "PanganFlow", regex=False)
    for col in ["ndcg_at_k", "precision_at_k", "recall_at_k", "lift_over_weighted"]:
        if col in metric_view.columns:
            metric_view[col] = metric_view[col].round(3)
    add_table(doc, metric_view, "Tabel II. Hasil evaluasi ranking provinsi prioritas.", font_size=6.8)
    add_figure(doc, FIGURE_DIR / "model_comparison.png", "Gambar 1. Perbandingan NDCG@10 antar model.", width=3.2)
    add_paragraph(
        doc,
        f"Model hybrid PanganFlow mencapai NDCG@10 sebesar {simple_metric(metrics, 'PanganFlow hybrid model', 'ndcg_at_k'):.3f}, sedangkan weighted index mencapai {simple_metric(metrics, 'Weighted priority index', 'ndcg_at_k'):.3f}. "
        "Interpretasinya, baseline transparan sudah sangat kuat karena selaras dengan definisi proxy, sedangkan Random Forest dipakai sebagai kalibrasi non-linear agar skor akhir tetap interpretabel."
    )
    ablation_view = ablation[["setting", "ndcg_at_k", "precision_at_k", "recall_at_k"]].copy()
    for col in ["ndcg_at_k", "precision_at_k", "recall_at_k"]:
        ablation_view[col] = ablation_view[col].round(3)
    add_table(doc, ablation_view, "Tabel III. Ablation study keluarga fitur.", font_size=7.0)
    add_figure(doc, FIGURE_DIR / "feature_importance.png", "Gambar 2. Feature importance model prioritas.", width=3.2)

    add_heading(doc, "VI. ANALISIS SPASIAL DAN REKOMENDASI FLOW")
    add_paragraph(
        doc,
        "Ranking provinsi digunakan sebagai sisi tujuan, sedangkan skor surplus digunakan sebagai sisi asal. Sistem tidak mengklaim rute operasional final; keluaran dipakai sebagai antrian review untuk melihat apakah terdapat selisih harga dan defisit tujuan yang cukup kuat dibanding penalti jarak."
    )
    flow_view = flows[["rank", "origin_province", "destination_province", "priority_score"]].head(6).copy()
    add_table(doc, flow_view, "Tabel IV. Top rekomendasi aliran beras.", font_size=6.6)
    for _, flow in flows.head(3).iterrows():
        add_paragraph(
            doc,
            f"Flow peringkat {int(flow['rank'])}: {flow['origin_province']} menuju {flow['destination_province']} memperoleh skor {float(flow['priority_score']):.1f}. "
            f"Alasan utama: {flow['top_reasons']}."
        )
    add_figure(doc, FIGURE_DIR / "pressure_centroid_map.png", "Gambar 3. Peta centroid tekanan beras nasional.", width=3.25)
    add_figure(doc, FIGURE_DIR / "flow_network.png", "Gambar 4. Visualisasi graph recommendation antarprovinsi.", width=3.25)

    add_heading(doc, "VII. CLUSTERING, FAIRNESS, DAN STUDI KASUS")
    cluster_view = clusters[["cluster_label", "province_count", "avg_deficit_score", "avg_surplus_score"]].copy()
    for col in ["avg_deficit_score", "avg_surplus_score"]:
        cluster_view[col] = cluster_view[col].round(3)
    add_table(doc, cluster_view, "Tabel V. Ringkasan cluster provinsi.", font_size=6.8)
    add_figure(doc, FIGURE_DIR / "cluster_pca.png", "Gambar 5. PCA cluster provinsi.", width=3.25)
    fairness_view = fairness[["region", "top10_count", "top10_share", "avg_priority_score"]].copy()
    for col in ["top10_share", "avg_priority_score"]:
        fairness_view[col] = fairness_view[col].round(3)
    add_table(doc, fairness_view, "Tabel VI. Pemeriksaan subgroup berdasarkan region.", font_size=6.8)
    case_view = cases[["province", "recommended_origin", "interpretation"]].copy()
    add_table(doc, case_view, "Tabel VII. Studi kasus provinsi prioritas.", font_size=6.4)
    add_paragraph(
        doc,
        f"Stability check dengan perturbasi skor menghasilkan rata-rata overlap Top-10 sebesar {float(stability['mean_top10_overlap'].iloc[0]):.2f}. "
        "Nilai ini menunjukkan ranking relatif stabil untuk presentasi awal, walaupun validasi lapangan tetap diperlukan sebelum digunakan sebagai dasar keputusan operasional."
    )

    add_heading(doc, "VIII. DASHBOARD DAN IMPLEMENTASI")
    add_paragraph(
        doc,
        "Dashboard Streamlit PanganFlow.ID dirancang sebagai alat presentasi dan eksplorasi. Tampilan utamanya terdiri atas peta tekanan beras, Top-K rekomendasi aliran, profil provinsi, feature explanations, tabel sumber data, dan panel validasi. "
        "Desain dashboard menempatkan rekomendasi sebagai prioritas review manusia agar tidak disalahpahami sebagai instruksi distribusi otomatis."
    )
    add_figure(doc, FIGURE_DIR / "dashboard_preview.png", "Gambar 6. Pratinjau dashboard PanganFlow.ID.", width=3.25)

    add_heading(doc, "IX. KETERBATASAN")
    add_bullets(
        doc,
        [
            "Neraca surplus-defisit memakai proxy tahunan; stok awal, cadangan, arus masuk, dan arus keluar belum dimodelkan penuh.",
            "Harga Bapanas bulanan lebih stabil untuk reproduksi, tetapi tidak sedetail harga harian PIHPS.",
            "Jarak centroid hanya proxy biaya logistik; moda transportasi, pelabuhan, gudang, dan kapasitas distribusi belum masuk.",
            "Model ranking mempelajari proxy prioritas, bukan ground-truth kebijakan resmi.",
        ],
    )

    add_heading(doc, "X. KESIMPULAN")
    add_paragraph(
        doc,
        "PanganFlow.ID menunjukkan bahwa data harga, produksi, konsumsi, dan spasial dapat digabung menjadi sistem penambangan data yang lebih kaya daripada prediksi harga tunggal. "
        "Output utamanya adalah ranking provinsi prioritas, cluster kondisi pangan, dan rekomendasi aliran surplus-defisit yang dapat dijelaskan. "
        "Dengan perluasan data stok dan logistik, sistem ini dapat dikembangkan menjadi early warning dan decision-support untuk mendukung kemandirian pangan nasional."
    )
    add_heading(doc, "Pengembangan Lanjutan", 2)
    add_paragraph(
        doc,
        "Tahap berikutnya adalah memasukkan stok awal, cadangan pemerintah, data distribusi masuk-keluar, kalender panen, dan indikator gangguan transportasi. "
        "Dengan variabel tersebut, skor flow dapat ditingkatkan dari prioritas review menjadi optimasi berbatas kapasitas yang mempertimbangkan biaya, waktu tempuh, dan kapasitas gudang."
    )
    add_paragraph(
        doc,
        "Dari sisi evaluasi, tim dapat menambah validasi pakar pangan, membandingkan rekomendasi dengan pola distribusi historis, serta melakukan backtesting pada periode kenaikan harga beras. "
        "Pengembangan multi-komoditas juga dapat dilakukan setelah tersedia data produksi dan konsumsi yang konsisten untuk cabai, bawang, telur, dan minyak goreng."
    )

    doc.add_page_break()
    add_heading(doc, "XI. REPRODUCIBILITY CHECKLIST")
    add_paragraph(
        doc,
        "Artefak dibuat agar dapat dijalankan ulang oleh tim maupun reviewer internal. Struktur proyek memisahkan data mentah, data terproses, tabel laporan, figur, source code, dashboard, dan dokumen final. "
        "Perintah publik yang disediakan adalah build dataset, train model, build report, dan menjalankan dashboard Streamlit."
    )
    add_bullets(
        doc,
        [
            "Dataset utama: data/processed/province_commodity_panel.csv.",
            "Rekomendasi utama: data/processed/flow_recommendations.csv.",
            "Evaluasi: reports/tables/metrics_summary.csv, feature_ablation.csv, ranking_stability.csv.",
            "Validasi: reports/tables/validation_checks.csv dan source_status.csv.",
            "Dokumen final: DOCX dan PDF dengan penamaan GEMASTIK placeholder ID_TIM dan NAMA_TIM.",
        ],
    )
    validation_view = validation[["check", "status", "value"]].tail(7).copy()
    add_table(doc, validation_view, "Tabel VIII. Ringkasan validasi artefak.", font_size=7.8)

    add_heading(doc, "DAFTAR PUSTAKA")
    refs = [
        "Bank Indonesia. Pusat Informasi Harga Pangan Strategis Nasional. https://www.bi.go.id/hargapangan/home/index/widget",
        "Bank Indonesia. User Manual Pusat Informasi Harga Pangan Strategis. https://www.bi.go.id/hargapangan/images/manual.pdf",
        "Badan Pangan Nasional. Rata-rata Harga Pangan Bulanan Tingkat Konsumen Provinsi. https://satudata.badanpangan.go.id/datasetpublications/tq2/harga-pangan-tingkat-konsumen-provinsi",
        "Badan Pusat Statistik. Produksi Padi dan Beras Menurut Provinsi. https://www.bps.go.id/id/statistics-table/3/ZDNaak0yODBUVTlGYW5sa2REUkVUVVY1YVZkbmR6MDkjMyMwMDAw/produksi-padi-sup-1-sup-dan-beras-menurut-provinsi.html?year=2024",
        "Badan Pusat Statistik. Pengeluaran untuk Konsumsi Penduduk Indonesia per Provinsi, Maret 2024. https://www.bps.go.id/id/publication/2024/10/18/dda144ad01aec46898795ccf/pengeluaran-untuk-konsumsi-penduduk-indonesia-per-provinsi--maret-2024.html/",
        "L. Breiman. Random Forests. Machine Learning, 2001. https://doi.org/10.1023/A:1010933404324",
        "F. T. Liu, K. M. Ting, and Z. H. Zhou. Isolation Forest. ICDM, 2008. https://doi.org/10.1109/ICDM.2008.17",
        "J. MacQueen. Some Methods for Classification and Analysis of Multivariate Observations. 1967. https://projecteuclid.org/euclid.bsmsp/1200512992",
    ]
    for idx, ref in enumerate(refs, start=1):
        p = doc.add_paragraph(style=style_name(doc, ["IEEEReferenceItem"], "Normal"))
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        r = p.add_run(f"[{idx}] {ref}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(8.4)

    doc.save(DOCX_PATH)


def export_pdf_with_word() -> bool:
    script = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{str(DOCX_PATH)}')
$doc.SaveAs([ref]'{str(PDF_PATH)}', [ref]17)
$doc.Close($false)
$word.Quit()
"""
    try:
        subprocess.run(["powershell", "-NoProfile", "-Command", script], check=True, capture_output=True, text=True, timeout=120)
        return PDF_PATH.exists()
    except Exception as exc:
        (REPORT_DIR / "pdf_export_error.log").write_text(f"Word COM export failed:\n{exc}\n", encoding="utf-8")
        return False


def find_render_docx_script() -> Path | None:
    base = Path.home() / ".codex" / "plugins" / "cache" / "openai-primary-runtime" / "documents"
    candidates = sorted(base.glob("*/skills/documents/render_docx.py"), reverse=True)
    return candidates[0] if candidates else None


def render_docx_pages() -> list[Path]:
    script = find_render_docx_script()
    if script is None:
        return []
    out_dir = REPORT_DIR / "rendered_docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            sys.executable,
            str(script),
            str(DOCX_PATH),
            "--output_dir",
            str(out_dir),
            "--renderer",
            "artifact-tool",
        ],
        check=True,
        capture_output=True,
        text=True,
        timeout=240,
    )

    def page_key(path: Path) -> int:
        match = re.search(r"page-(\d+)", path.stem)
        return int(match.group(1)) if match else 0

    return sorted(out_dir.glob("page-*.png"), key=page_key)


def export_pdf_from_rendered_pages() -> bool:
    try:
        pages = render_docx_pages()
        if not pages:
            return False
        images = []
        for page in pages:
            img = Image.open(page).convert("RGB")
            img.thumbnail((1240, 1754))
            canvas = Image.new("RGB", (1240, 1754), "white")
            canvas.paste(img, ((1240 - img.width) // 2, (1754 - img.height) // 2))
            images.append(canvas)
        images[0].save(PDF_PATH, "PDF", resolution=150.0, save_all=True, append_images=images[1:], quality=86)
        return PDF_PATH.exists()
    except Exception as exc:
        with (REPORT_DIR / "pdf_export_error.log").open("a", encoding="utf-8") as fh:
            fh.write(f"\nRendered-page PDF fallback failed:\n{exc}\n")
        return False


def main() -> None:
    parser = argparse.ArgumentParser(description="Build PanganFlow.ID formal GEMASTIK report.")
    parser.add_argument("--skip-pdf", action="store_true")
    args = parser.parse_args()
    write_report()
    ok = False
    if not args.skip_pdf:
        ok = export_pdf_with_word() or export_pdf_from_rendered_pages()
    print(f"DOCX: {DOCX_PATH}")
    if not args.skip_pdf:
        print(f"PDF: {PDF_PATH if ok else 'export failed; see reports/pdf_export_error.log'}")


if __name__ == "__main__":
    main()
