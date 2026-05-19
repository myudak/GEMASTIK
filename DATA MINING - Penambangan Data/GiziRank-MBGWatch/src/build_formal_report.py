from __future__ import annotations

import json
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
TABLES = REPORTS / "tables"
FIGURES = REPORTS / "figures"
DATA = ROOT / "data" / "processed"

DOCX_NAME = "GEMASTIK XVIII Penambangan Data - ID_TIM - NAMA_TIM - GiziRank.docx"
PDF_NAME = "GEMASTIK XVIII Penambangan Data - ID_TIM - NAMA_TIM - GiziRank.pdf"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_run(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str, style: str | None = None, align: int | None = None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    style_run(run, 11)
    return paragraph


def add_heading(doc: Document, number: str, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(f"{number}. {title}")
    style_run(run, 14, True, "12343B")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    style_run(run, 9, False, "465A62")
    run.italic = True


def add_table(doc: Document, df: pd.DataFrame, caption: str, widths: list[float] | None = None, max_rows: int | None = None) -> None:
    if max_rows is not None:
        df = df.head(max_rows)
    df = df.copy()
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(caption)
    style_run(run, 9, True, "12343B")
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.autofit = True
    for i, col in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(col)
        set_cell_shading(cell, "E9F1ED")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                style_run(run, 8, True, "12343B")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float):
                text = f"{value:.3f}" if abs(value) < 1000 else f"{value:,.0f}"
            else:
                text = str(value)
            cells[i].text = text
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    style_run(run, 8)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def configure_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(11)
    return doc


def export_pdf(docx_path: Path, pdf_path: Path) -> bool:
    script = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{str(docx_path.resolve())}')
$doc.SaveAs([ref] '{str(pdf_path.resolve())}', [ref] 17)
$doc.Close($false)
$word.Quit()
"""
    try:
        subprocess.run(["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script], check=True, capture_output=True, text=True)
        return True
    except Exception as exc:  # noqa: BLE001
        print(f"PDF export failed: {exc}")
        return False


def main() -> int:
    REPORTS.mkdir(parents=True, exist_ok=True)
    panel = pd.read_csv(DATA / "district_panel.csv")
    ranking = pd.read_csv(TABLES / "ranking_snapshot.csv")
    metrics = pd.read_csv(TABLES / "metrics_summary.csv")
    importance = pd.read_csv(TABLES / "feature_importance.csv")
    ablation = pd.read_csv(TABLES / "feature_ablation.csv")
    fairness = pd.read_csv(TABLES / "subgroup_fairness.csv")
    stability = pd.read_csv(TABLES / "ranking_stability.csv")
    summary = pd.read_csv(TABLES / "dataset_summary.csv")
    source_status = pd.read_csv(TABLES / "source_status.csv")
    cases = pd.read_csv(TABLES / "case_studies.csv")
    model_card = json.loads((TABLES / "model_card.json").read_text(encoding="utf-8"))

    doc = configure_doc()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("GiziRank: Data Mining untuk Prioritas Monitoring dan Dukungan Program Makan Bergizi Gratis")
    style_run(run, 17, True, "12343B")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(subtitle.add_run("GEMASTIK XVIII - Penambangan Data"), 11, True, "465A62")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(meta.add_run("<ID Tim> - <Nama Tim> - <Universitas>\n<Anggota 1>, <Anggota 2>, <Anggota 3>"), 10, False, "465A62")

    add_heading(doc, "0", "Abstrak")
    add_paragraph(
        doc,
        "Program Makan Bergizi Gratis (MBG) membutuhkan mekanisme prioritas monitoring yang terukur karena skala penerima dan unit layanan terus bertambah. "
        "Makalah ini mengusulkan GiziRank, sistem pemeringkatan kabupaten/kota untuk membantu menentukan wilayah yang perlu diprioritaskan dalam dukungan kualitas layanan, perluasan cakupan, dan pemantauan berbasis data. "
        "Dataset dibangun dari sumber publik resmi, yaitu Dashboard MBG Kemendikdasmen, direktori SPPG Operasional BGN, Indeks Ketahanan Pangan Bapanas, harga pangan Bapanas, dan tabel kemiskinan BPS bila endpoint tersedia. "
        "Metode yang digunakan terdiri dari indeks risiko transparan, model ranking berbasis ensemble decision-stump, evaluasi Precision@K, Recall@K, NDCG@K, ablation study, serta analisis TransferRisk antarkelompok wilayah. "
        "Hasil awal menunjukkan model ranking mampu menghasilkan queue prioritas yang dapat dijelaskan melalui faktor kebutuhan gizi, gap cakupan, kerentanan pangan, tekanan sosial-ekonomi, dan tekanan harga. "
        "GiziRank tidak mengklaim penilaian individu atau keamanan dapur, melainkan menyediakan alat bantu prioritas untuk monitoring program publik.",
    )
    add_paragraph(doc, "Kata kunci: data mining, MBG, learning to rank, monitoring prioritas, ketahanan pangan, explainable AI.")

    add_heading(doc, "1", "Pendahuluan")
    add_paragraph(
        doc,
        "Tema GEMASTIK 2025 menekankan peningkatan TIK menuju kemandirian bangsa. Dalam konteks tersebut, MBG relevan karena menyentuh kualitas sumber daya manusia, ketahanan pangan, dan tata kelola layanan publik. "
        "Masalah praktisnya bukan sekadar menampilkan dashboard, tetapi membuat daftar prioritas yang dapat dipertanggungjawabkan ketika kapasitas monitoring terbatas.",
    )
    add_paragraph(
        doc,
        "Tujuan penelitian adalah membangun pipeline data mining yang mengintegrasikan sinyal kebutuhan, cakupan, pangan, dan sosial-ekonomi untuk menghasilkan ranking kabupaten/kota. "
        "Manfaatnya adalah membantu pengambil keputusan membaca wilayah prioritas dengan metrik Top-K dan alasan yang dapat dijelaskan. Batasan utama: unit analisis adalah kabupaten/kota, label adalah proxy prioritas, dan keluaran tidak digunakan untuk menilai individu, rumah tangga, atau SPPG tertentu.",
    )

    add_heading(doc, "2", "Kajian Terkait")
    add_paragraph(
        doc,
        "Proyek ini mengambil bentuk masalah learning-to-rank dan recommender system, serupa pola tugas KDD yang menilai kemampuan model menyusun prioritas item berdasarkan relevansi. "
        "Dalam domain kebijakan publik, pendekatan Top-K lebih cocok daripada klasifikasi biasa karena keputusan lapangan sering berbentuk alokasi sumber daya terbatas. "
        "Evaluasi menggunakan NDCG@K, Precision@K, dan Recall@K mengikuti praktik information retrieval dan ranking. "
        "Penggunaan model pohon dipilih karena kuat pada data tabular, sedangkan feature importance dipakai untuk menjaga keterjelasan argumen kepada juri dan pemangku kepentingan.",
    )

    add_heading(doc, "3", "Dataset dan Formulasi Masalah")
    compact_source = source_status[["source", "status", "rows", "notes"]].copy()
    add_table(doc, compact_source, "Tabel 1. Status sumber data publik yang digunakan.", max_rows=8)
    summary_pivot = summary.rename(columns={"metric": "Metrik", "value": "Nilai"})
    add_table(doc, summary_pivot, "Tabel 2. Ringkasan dataset terintegrasi.", max_rows=12)
    add_paragraph(
        doc,
        "Setiap baris dataset merepresentasikan satu kabupaten/kota pada periode snapshot MBG terbaru. "
        "Sinyal dukungan diwakili jumlah penerima MBG, satuan pendidikan, dan jumlah SPPG operasional yang berhasil diagregasikan. "
        "Sinyal kebutuhan diwakili rasio kondisi khusus peserta, kerentanan pangan IKP, tekanan harga, dan kemiskinan bila tabel BPS dapat diakses.",
    )
    add_table(
        doc,
        pd.DataFrame(
            [
                ["nutrition_need", "Skor 0-1 dari kondisi khusus peserta, alergi/intoleransi, dan kerentanan pangan"],
                ["coverage_gap", "Selisih ternormalisasi antara need pressure dan support reach"],
                ["priority_label", "Top 20% reference_priority_score sebagai proxy prioritas monitoring"],
                ["final_priority_score", "Gabungan skor model ranking dan skor referensi transparan"],
            ],
            columns=["Variabel", "Definisi"],
        ),
        "Tabel 3. Definisi target dan skor utama.",
    )

    add_heading(doc, "4", "Metodologi")
    add_paragraph(
        doc,
        "Pipeline dimulai dari pengambilan data resmi, normalisasi nama provinsi/kabupaten, agregasi SPPG ke level kabupaten/kota, penggabungan sumber lintas lembaga, dan pembentukan fitur. "
        "Skor referensi dihitung dari kombinasi nutrition need, coverage gap, food vulnerability, socioeconomic stress, dan price stress. "
        "Model utama adalah ensemble decision-stump dengan objective logistik sederhana. Implementasi ini sengaja ringan agar dapat direplikasi tanpa dependensi ML berat.",
    )
    add_paragraph(
        doc,
        f"Model final menggunakan {model_card['n_estimators_fit']} stump dari {len(model_card['feature_names'])} fitur. "
        "Validasi dilakukan dengan holdout berbasis provinsi agar evaluasi tidak hanya mengukur kemampuan menghafal distrik dalam provinsi yang sama.",
    )

    add_heading(doc, "5", "Eksperimen dan Pengujian")
    add_table(doc, metrics.round(4), "Tabel 4. Perbandingan baseline dan model utama pada province holdout.")
    add_image(doc, FIGURES / "model_comparison.png", "Gambar 1. Perbandingan NDCG@K antar model.", width=6.2)
    add_table(doc, ablation.round(4), "Tabel 5. Ablation study menurut keluarga fitur.")
    add_paragraph(
        doc,
        "Baseline acak menunjukkan batas bawah performa ranking. Weighted index baseline menguji apakah formula transparan saja sudah cukup. "
        "StumpBoost Ranker dipakai untuk menangkap interaksi non-linear sederhana antarfitur tanpa mengorbankan replikasi.",
    )

    add_heading(doc, "6", "Hasil dan Analisis")
    top10 = ranking[["rank", "district", "province", "final_priority_score", "top_reasons"]].head(10).copy()
    top10["alasan_ringkas"] = top10["top_reasons"].map(lambda value: "; ".join(str(value).split("; ")[:2]))
    top10 = top10.drop(columns=["top_reasons"])
    add_table(doc, top10, "Tabel 6. Sepuluh kabupaten/kota prioritas tertinggi.", widths=[0.45, 1.15, 1.05, 0.8, 2.7])
    add_image(doc, FIGURES / "top_priority_districts.png", "Gambar 2. Ranking prioritas tertinggi.", width=6.2)
    add_image(doc, FIGURES / "risk_quadrant.png", "Gambar 3. Kuadran risiko berbasis nutrition need dan coverage gap.", width=6.2)
    add_paragraph(
        doc,
        "Wilayah pada kuadran kanan-atas memiliki kombinasi kebutuhan tinggi dan gap dukungan tinggi. "
        "Bagian ini paling berguna untuk prioritas monitoring karena menunjukkan wilayah yang bukan hanya rentan, tetapi juga relatif belum seimbang dengan sinyal dukungan yang tersedia.",
    )

    add_heading(doc, "7", "Explainability dan TransferRisk")
    add_table(doc, importance[["feature", "family", "importance"]].head(10).round(4), "Tabel 7. Feature importance model ranking.", max_rows=10)
    add_image(doc, FIGURES / "feature_importance.png", "Gambar 4. Kontribusi fitur utama.", width=6.2)
    doc.add_page_break()
    add_table(doc, fairness.round(4), "Tabel 8. Analisis TransferRisk dan subgroup.", max_rows=8)
    add_image(doc, FIGURES / "subgroup_top10_share.png", "Gambar 5. Proporsi Top-K per kelompok wilayah.", width=6.2)
    add_paragraph(
        doc,
        f"Stabilitas Top-20 diuji dengan perturbasi skor kecil. Rata-rata overlap Top-20 adalah {stability['bootstrap_top20_overlap_mean'].iloc[0]:.2f}, "
        f"dengan persentil ke-10 sebesar {stability['bootstrap_top20_overlap_p10'].iloc[0]:.2f}. "
        "Nilai ini menunjukkan apakah queue prioritas sensitif terhadap perubahan kecil pada skor model.",
    )

    doc.add_page_break()
    add_heading(doc, "8", "Dashboard MBGWatch")
    add_image(doc, FIGURES / "dashboard_preview.png", "Gambar 6. Preview dashboard MBGWatch untuk eksplorasi ranking.", width=6.2)
    add_paragraph(
        doc,
        "Dashboard menampilkan metrik utama, queue Top-K, profil kabupaten/kota, alasan prioritas, dan ringkasan provinsi. "
        "Peran dashboard adalah lapisan pembacaan hasil data mining; kontribusi ilmiah tetap berada pada formulasi ranking, pipeline dataset, evaluasi Top-K, dan analisis risiko transfer.",
    )

    doc.add_page_break()
    add_heading(doc, "9", "Studi Kasus dan Keterbatasan")
    case_table = cases[["rank", "district", "province", "case_note"]].copy()
    add_table(doc, case_table, "Tabel 9. Studi kasus Top-5.", widths=[0.45, 1.1, 1.1, 3.9], max_rows=5)
    add_paragraph(
        doc,
        "Keterbatasan utama adalah label prioritas yang masih berbasis proxy, bukan hasil anotasi ahli. "
        "Sebagian sumber publik juga berbeda granularitas dan periode waktunya. Jika tim memperoleh snapshot stunting atau total eligible students yang lebih lengkap, pipeline sudah menyediakan titik masuk manual untuk memperkuat definisi label dan coverage gap.",
    )

    add_heading(doc, "10", "Kesimpulan")
    add_paragraph(
        doc,
        "GiziRank menunjukkan bahwa masalah MBG dapat diformulasikan sebagai tugas data mining yang jelas: ranking prioritas monitoring kabupaten/kota. "
        "Dengan sumber publik resmi, model ringan, metrik Top-K, dan dashboard interpretatif, solusi ini lebih kuat daripada dashboard statis karena menghasilkan queue aksi yang terukur. "
        "Pengembangan berikutnya adalah validasi ahli, integrasi data stunting tabular terbaru, dan evaluasi prospektif terhadap tindak lanjut monitoring lapangan.",
    )

    add_heading(doc, "11", "Daftar Pustaka")
    refs = [
        "[1] Badan Gizi Nasional. SPPG Operasional. https://www.bgn.go.id/operasional-sppg",
        "[2] Kemendikdasmen. Dashboard MBG. https://mbg.pdm.kemendikdasmen.go.id/portal",
        "[3] Dashboard TP2S. Dashboard Pemantauan Terpadu Percepatan Pencegahan Stunting. https://dashboard.stunting.go.id/",
        "[4] Badan Pangan Nasional. Indeks Ketahanan Pangan Kabupaten/Kota 2024. https://data.badanpangan.go.id/datasetpublications/frq/ikp-kab-kota-2024",
        "[5] Badan Pangan Nasional. Harga Pangan Tingkat Konsumen Provinsi. https://satudata.badanpangan.go.id/datasetpublications/tq2/harga-pangan-tingkat-konsumen-provinsi",
        "[6] Badan Pusat Statistik. Persentase Penduduk Miskin menurut Kabupaten/Kota. https://www.bps.go.id/id/statistics-table/2/NjIxIzI=/persentase-penduduk-miskin--p0--menurut-kabupaten-kota.html",
        "[7] Davis, J. and Goadrich, M. The Relationship Between Precision-Recall and ROC Curves. ICML, 2006.",
        "[8] Friedman, J. Greedy Function Approximation: A Gradient Boosting Machine. Annals of Statistics, 2001.",
        "[9] Breiman, L. Random Forests. Machine Learning, 2001.",
    ]
    for ref in refs:
        add_paragraph(doc, ref)

    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    add_heading(doc, "Lampiran A", "Reproducibility Checklist")
    checklist = pd.DataFrame(
        [
            ["Dataset builder", "python src/build_dataset.py"],
            ["Model training", "python src/train_gizirank.py"],
            ["Report generation", "python src/build_formal_report.py"],
            ["Dashboard", "streamlit run dashboard/app.py"],
            ["Manual review", "reports/tables/unmatched_districts_for_manual_review.csv"],
            ["Static preview", "dashboard/static_preview.html"],
        ],
        columns=["Artefak", "Lokasi/Perintah"],
    )
    add_table(doc, checklist, "Tabel A1. Reproducibility checklist.")

    docx_path = REPORTS / DOCX_NAME
    pdf_path = REPORTS / PDF_NAME
    doc.save(docx_path)
    exported = export_pdf(docx_path, pdf_path)
    print(f"DOCX written: {docx_path}")
    print(f"PDF written: {pdf_path if exported else 'not exported'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
